from __future__ import annotations

import hashlib
import json
import os
import threading
import uuid
from io import BytesIO
from pathlib import Path
from typing import Iterable, List, Optional, Tuple

import easyocr
import numpy as np
import ollama
from docx import Document as DocxDocument
from langchain_community.document_loaders import PDFPlumberLoader
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_core.prompts import PromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter
from PIL import Image

from services.notebook_storage import NotebookStorage
from services.graph_rag_service import GraphRAGService


class LocalOllamaEmbeddings(Embeddings):
    def __init__(self, model: str = "nomic-embed-text") -> None:
        host = os.getenv("OLLAMA_HOST")
        self.client = ollama.Client(host=host) if host else ollama.Client()
        self.model = model

    def _extract_embeddings(self, response) -> List[List[float]]:
        if hasattr(response, "embeddings"):
            return response.embeddings
        if isinstance(response, dict) and "embeddings" in response:
            return response["embeddings"]
        raise ValueError("Unexpected embedding response format from Ollama")

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        if not texts:
            return []
        response = self.client.embed(model=self.model, input=texts)
        return self._extract_embeddings(response)

    def embed_query(self, text: str) -> List[float]:
        return self.embed_documents([text])[0]


class RAGService:
    SUPPORTED_FILE_EXTENSIONS = {".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff", ".tif"}

    def __init__(self, storage: Optional[NotebookStorage] = None) -> None:
        self.storage = storage or NotebookStorage()
        self.embedder = LocalOllamaEmbeddings()
        self._ocr_reader: Optional[easyocr.Reader] = None
        #tạo graph 
        self.GRS = GraphRAGService(self.storage)
        self._index_lock = threading.Lock()
        

    def _get_ocr_reader(self) -> easyocr.Reader:
        if self._ocr_reader is None:
            self._ocr_reader = easyocr.Reader(["en", "vi"], gpu=False)
        return self._ocr_reader

    def _ollama_client(self) -> ollama.Client:
        host = os.getenv("OLLAMA_HOST")
        return ollama.Client(host=host) if host else ollama.Client()

    def _hash_bytes(self, content: bytes) -> str:
        return hashlib.sha256(content).hexdigest()

    def extract_text_from_image_bytes(self, image_bytes: bytes) -> str:
        image = Image.open(BytesIO(image_bytes)).convert("RGB")
        image_array = np.array(image)
        lines = self._get_ocr_reader().readtext(image_array, detail=0)
        return "\n".join(lines).strip()

    def _build_system_prompt(self, answer_language: str, chat_settings: Optional[dict[str, str]] = None) -> str:
        settings = chat_settings or {}
        response_length = settings.get("response_length", "medium")
        roleplay = (settings.get("roleplay") or "").strip()
        mode = settings.get("mode", "normal")

        if answer_language == "vi":
            length_instruction = {
                "short": "Hãy trả lời cực kỳ ngắn gọn dưới 3 câu.",
                "medium": "Hãy trả lời vừa phải, khoảng 3-5 câu.",
                "long": "Hãy trả lời chi tiết hơn, khoảng 6-10 câu, nhưng vẫn tập trung vào thông tin cần thiết.",
            }[response_length]
            mode_instruction = {
                "normal": "Trả lời bình thường, rõ ràng và chính xác.",
                "study_guide": "Ưu tiên cấu trúc như một tài liệu hướng dẫn học tập: giải thích khái niệm, chia bước, thêm mẹo ghi nhớ khi phù hợp.",
                "critical_thinking": "Ưu tiên phân tích phản biện: nêu giả định, điểm yếu, rủi ro, và câu hỏi cần kiểm chứng.",
            }[mode]
            base = [
                "Bạn là trợ lý Local RAG. Chỉ trả lời dựa trên ngữ cảnh được cung cấp.",
                "Bạn được cung cấp ngữ cảnh từ nhiều nguồn tài liệu khác nhau. Khi trả lời, bắt buộc phải TRÍCH DẪN RÕ RÀNG thông tin lấy từ tài liệu nào. TUYỆT ĐỐI KHÔNG trộn lẫn, tự ý kết nối các sự kiện của tài liệu này sang tài liệu khác nếu chúng không liên quan trong văn bản.",
                "Khi sử dụng thông tin từ nhiều tài liệu, hãy giữ ranh giới giữa các nguồn và chỉ kết luận khi các nguồn thực sự hỗ trợ cho cùng một điểm.",
                length_instruction,
                mode_instruction,
                "Nếu ngữ cảnh không đủ, hãy nói rõ là không có đủ thông tin.",
                "Hãy dùng ký hiệu trích dẫn [1], [2], [3] tương ứng với các khối ngữ cảnh được cung cấp.",
                "Luôn trả lời bằng tiếng Việt, trừ khi người dùng yêu cầu khác.",
            ]
        else:
            length_instruction = {
                "short": "Keep the answer extremely concise, under 3 sentences.",
                "medium": "Keep the answer moderate, around 3-5 sentences.",
                "long": "Provide a deeper answer, around 6-10 sentences, while staying focused on the necessary information.",
            }[response_length]
            mode_instruction = {
                "normal": "Answer normally, clearly, and accurately.",
                "study_guide": "Prefer a study-guide style: explain concepts, break steps down, and add memory aids when relevant.",
                "critical_thinking": "Prefer critical analysis: surface assumptions, weaknesses, risks, and questions that need verification.",
            }[mode]
            base = [
                "You are a Local RAG assistant. Answer only from the provided context.",
                "You are given context from multiple different documents. When answering, you must clearly cite which document each fact comes from. DO NOT mix, merge, or invent connections between unrelated facts from different documents.",
                "When using multiple sources, keep them separate and only combine them when the text itself supports that relationship.",
                length_instruction,
                mode_instruction,
                "If the context is insufficient, say so clearly.",
                "Use citation markers like [1], [2], [3] that match the numbered context blocks provided.",
                "Always answer in English unless the user explicitly asks for another language.",
            ]

        if roleplay:
            if answer_language == "vi":
                base.insert(3, f"Bạn đang đóng vai: {roleplay}.")
            else:
                base.insert(3, f"You are roleplaying as: {roleplay}.")

        return "\n".join(base)

    def _build_prompt_template(self, answer_language: str, chat_settings: Optional[dict[str, str]] = None) -> PromptTemplate:
        system_prompt = self._build_system_prompt(answer_language, chat_settings)
        if answer_language == "vi":
            template = (
                f"{system_prompt}\n\n"
                "Ngu canh:\n{context}\n\n"
                "Cau hoi: {question}\n\n"
                "Tra loi ngan gon 3-5 cau, uu tien thong tin chinh xac tu nguon."
            )
        else:
            template = (
                f"{system_prompt}\n\n"
                "Context:\n{context}\n\n"
                "Question: {question}\n\n"
                "Keep the answer concise in 3-5 sentences and grounded in the sources."
            )
        return PromptTemplate(template=template, input_variables=["context", "question"])

    def _is_vietnamese(self, text: str) -> bool:
        chars = "aadeioouuyáàảãạăắằẳẵặâấầẩẫậđéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵ"
        lowered = text.lower()
        if any(c in lowered for c in chars):
            return True

        vi_keywords = {
            "toi", "ban", "la", "khong", "duoc", "cua", "trong", "cho", "voi",
            "nhung", "mot", "nhieu", "tai", "sao", "the", "nao", "bao", "nhieu",
            "cau", "hoi", "nguon", "tai", "lieu",
        }
        tokens = {token.strip(".,!?;:()[]{}\"'") for token in lowered.split() if token.strip()}
        return len(tokens.intersection(vi_keywords)) >= 2

    def _detect_answer_language(self, question: str) -> str:
        return "vi" if self._is_vietnamese(question) else "en"

    def _should_rewrite_query(self, question: str) -> bool:
        lowered = question.lower()
        follow_up_markers = {
            "it", "this", "that", "they", "them", "he", "she",
            "no", "do", "day", "kia", "tren", "duoi", "tiep",
        }
        tokens = {token.strip(".,!?;:()[]{}\"'") for token in lowered.split() if token.strip()}
        return len(tokens) <= 12 and len(tokens.intersection(follow_up_markers)) > 0

    def _parse_file_to_documents(self, file_path: Path, source_name: str, file_hash: str) -> Tuple[List[Document], int]:
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            docs = PDFPlumberLoader(str(file_path)).load()
            for d in docs:
                page_num = int(d.metadata.get("page", 1))
                d.metadata.update({"source_name": source_name, "page_number": page_num, "file_hash": file_hash})
            return docs, len(docs)

        if suffix == ".docx":
            docx = DocxDocument(str(file_path))
            text = "\n".join([p.text for p in docx.paragraphs if p.text])
            docs = [Document(page_content=text, metadata={"source_name": source_name, "page_number": 1, "file_hash": file_hash})]
            return docs, 1

        if suffix == ".txt":
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            docs = [Document(page_content=text, metadata={"source_name": source_name, "page_number": 1, "file_hash": file_hash})]
            return docs, 1

        if suffix in {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff", ".tif"}:
            image = Image.open(file_path).convert("RGB")
            image_array = np.array(image)
            lines = self._get_ocr_reader().readtext(image_array, detail=0)
            text = "\n".join(lines).strip()
            docs = [Document(page_content=text, metadata={"source_name": source_name, "page_number": 1, "file_hash": file_hash})]
            return docs, 1

        raise ValueError(f"Unsupported file type: {suffix}")

    def save_upload(self, notebook_id: str, filename: str, content: bytes) -> Optional[dict]:
        suffix = Path(filename).suffix.lower()
        if suffix not in self.SUPPORTED_FILE_EXTENSIONS:
            return None

        dirs = self.storage.get_notebook_dirs(notebook_id)
        dirs["uploads"].mkdir(parents=True, exist_ok=True)

        file_hash = self._hash_bytes(content)
        disk_name = f"{file_hash}{suffix}"
        upload_path = dirs["uploads"] / disk_name
        upload_path.write_bytes(content)

        document_id = uuid.uuid4().hex
        self.storage.upsert_source_document(
            document_id=document_id,
            notebook_id=notebook_id,
            source_name=filename,
            source_type=suffix.lstrip("."),
            file_hash=file_hash,
            upload_path=str(upload_path),
            page_count=None,
            file_size_bytes=len(content),
        )
        return {
            "document_id": document_id,
            "source_name": filename,
            "source_type": suffix.lstrip("."),
            "file_hash": file_hash,
            "upload_path": str(upload_path),
            "file_size_bytes": len(content),
        }

    def save_web_source(self, notebook_id: str, source_url: str, extracted_text: str) -> dict:
        if not extracted_text.strip():
            raise ValueError("No textual content extracted from URL")

        dirs = self.storage.get_notebook_dirs(notebook_id)
        dirs["uploads"].mkdir(parents=True, exist_ok=True)

        content = extracted_text.strip().encode("utf-8")
        file_hash = self._hash_bytes(content)
        disk_name = f"{file_hash}.txt"
        upload_path = dirs["uploads"] / disk_name
        upload_path.write_bytes(content)

        document_id = uuid.uuid4().hex
        self.storage.upsert_source_document(
            document_id=document_id,
            notebook_id=notebook_id,
            source_name=source_url,
            source_type="web_link",
            file_hash=file_hash,
            upload_path=str(upload_path),
            page_count=1,
            file_size_bytes=len(content),
        )
        return {
            "document_id": document_id,
            "source_name": source_url,
            "source_type": "web_link",
            "file_hash": file_hash,
            "upload_path": str(upload_path),
            "file_size_bytes": len(content),
        }

    def delete_source(self, document_id: str) -> dict:
        source = self.storage.get_source_document_by_id(document_id)
        if source is None:
            raise ValueError("Source not found")

        notebook_id = str(source["notebook_id"])
        upload_path = Path(str(source["upload_path"]))

        deleted = self.storage.delete_source_document(document_id=document_id)
        if deleted is None:
            raise ValueError("Source not found")

        if upload_path.exists() and upload_path.is_file():
            upload_path.unlink(missing_ok=True)

        remaining_sources = self.storage.list_source_documents(notebook_id)
        if remaining_sources:
            self.index_notebook(
                notebook_id=notebook_id,
                source_ids=None,
                chunk_size=1000,
                chunk_overlap=100,
            )
        else:
            self.storage.clear_index_data(notebook_id)
            dirs = self.storage.get_notebook_dirs(notebook_id)
            index_dir = dirs["vector_db"]
            for filename in ("index.faiss", "index.pkl"):
                file_path = index_dir / filename
                if file_path.exists() and file_path.is_file():
                    file_path.unlink(missing_ok=True)

        return {
            "notebook_id": notebook_id,
            "document_id": str(source["document_id"]),
            "source_name": str(source["source_name"]),
            "source_type": str(source["source_type"]),
        }

    def _load_or_create_vector_store(self, notebook_id: str) -> FAISS:
        dirs = self.storage.get_notebook_dirs(notebook_id)
        index_dir = dirs["vector_db"]
        index_file = index_dir / "index.faiss"

        if index_file.exists():
            return FAISS.load_local(
                folder_path=str(index_dir),
                embeddings=self.embedder,
                allow_dangerous_deserialization=True,
            )

        # Create a new FAISS store with a temporary seed document, then clear content.
        seed = Document(page_content="seed", metadata={"source_name": "seed", "page_number": 0})
        vector_store = FAISS.from_documents([seed], self.embedder)
        vector_store.delete(ids=list(vector_store.index_to_docstore_id.values()))
        return vector_store

    def index_notebook(
        self,
        notebook_id: str,
        source_ids: Optional[List[str]],
        chunk_size: int,
        chunk_overlap: int,
    ) -> dict:
        if not self.storage.notebook_exists(notebook_id):
            raise ValueError("Notebook not found")

        all_sources = self.storage.list_source_documents(notebook_id)
        if source_ids:
            source_id_set = set(source_ids)
            selected_sources = [row for row in all_sources if row["document_id"] in source_id_set]
        else:
            selected_sources = list(all_sources)

        if not selected_sources:
            raise ValueError("No source documents found to index")

        splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        vector_store = self._load_or_create_vector_store(notebook_id)

        self.storage.clear_index_data(notebook_id)

        total_chunks = 0
        total_pages = 0
        current_vector_size = len(vector_store.index_to_docstore_id)

        graph_chunks = []  # Collect chunks for graph building

        for src in selected_sources:
            source_docs, page_count = self._parse_file_to_documents(
                file_path=Path(src["upload_path"]),
                source_name=src["source_name"],
                file_hash=src["file_hash"],
            )
            total_pages += page_count

            chunks = splitter.split_documents(source_docs)
            if not chunks:
                continue

            # Build graph IMMEDIATELY after creating chunks (before database saves)
            graph_chunks.extend(chunks)

            chunk_ids: List[str] = []
            for idx, chunk in enumerate(chunks):
                chunk_id = uuid.uuid4().hex
                chunk.metadata.update(
                    {
                        "notebook_id": notebook_id,
                        "document_id": src["document_id"],
                        "chunk_index": idx,
                    }
                )
                page_number = chunk.metadata.get("page_number") or chunk.metadata.get("page") or 1
                page_number = int(page_number)
                chunk_ids.append(chunk_id)

                self.storage.insert_chunk(
                    chunk_id=chunk_id,
                    notebook_id=notebook_id,
                    document_id=src["document_id"],
                    chunk_index=idx,
                    page_number=page_number,
                    text_content=chunk.page_content,
                    metadata_json=json.dumps(chunk.metadata, ensure_ascii=True),
                )

            vector_store.add_documents(documents=chunks, ids=chunk_ids)

            for i, chunk_id in enumerate(chunk_ids):
                vector_id = current_vector_size + i
                page_number = int(chunks[i].metadata.get("page_number") or chunks[i].metadata.get("page") or 1)
                self.storage.upsert_vector_entry(
                    notebook_id=notebook_id,
                    faiss_vector_id=vector_id,
                    chunk_id=chunk_id,
                    source_name=src["source_name"],
                    page_number=page_number,
                )

            current_vector_size += len(chunk_ids)
            total_chunks += len(chunk_ids)

        # Build graph once after collecting all chunks (single call)
        print(f"[RAG_INDEX] Collected {len(graph_chunks)} total chunks for graph building")
        if graph_chunks:
            try:
                print(f"[RAG_INDEX] Calling build_graph_from_chunks with {len(graph_chunks)} chunks...")
                self.GRS.build_graph_from_chunks(graph_chunks)
                print(f"[RAG_INDEX] ✅ Graph building completed successfully")
            except Exception as e:
                print(f"[RAG_INDEX ERROR] GraphRAG indexing failed: {type(e).__name__}: {str(e)}")
                import traceback
                print(f"[RAG_INDEX ERROR] Traceback:\n{traceback.format_exc()}")
        else:
            print(f"[RAG_INDEX WARNING] No graph chunks collected - GraphRAG will be empty!")
            

        dirs = self.storage.get_notebook_dirs(notebook_id)
        dirs["vector_db"].mkdir(parents=True, exist_ok=True)
        vector_store.save_local(str(dirs["vector_db"]))

        return {
            "notebook_id": notebook_id,
            "indexed_sources": len(selected_sources),
            "total_pages": total_pages,
            "total_chunks": total_chunks,
            "vector_path": str(dirs["vector_db"]),
        }

    def _rewrite_query(self, notebook_id: str, session_id: str, question: str) -> str:
        history = self.storage.get_recent_chat_messages(notebook_id=notebook_id, session_id=session_id, limit=10)
        if not history:
            return question

        paired_history: list[tuple[str, str]] = []
        pending_user: str | None = None
        for item in history:
            role = str(item['role'])
            content = str(item['content']).strip()
            if not content:
                continue

            if role == 'user':
                pending_user = content
                continue

            if role == 'assistant' and pending_user:
                paired_history.append((pending_user, content))
                pending_user = None

        if not paired_history and history:
            history_lines = [f"{item['role']}: {item['content']}" for item in history]
        else:
            recent_pairs = paired_history[-5:]
            history_lines = []
            for index, (user_text, assistant_text) in enumerate(recent_pairs, start=1):
                history_lines.append(f"Pair {index} - User: {user_text}")
                history_lines.append(f"Pair {index} - Assistant: {assistant_text}")

        if self._is_vietnamese(question):
            prompt = (
                "Dua vao lich su hoi thoai sau, hay viet lai cau hoi cuoi cung cua nguoi dung "
                "thanh mot cau hoi doc lap, day du chu ngu, vi ngu va tu khoa de tim kiem tai lieu. "
                "Chi tra ve DUY NHAT cau hoi da viet lai, khong giai thich them.\n\n"
                f"Lich su hoi thoai:\n{os.linesep.join(history_lines)}\n\n"
                f"Cau hoi goc: {question}\n"
                "Cau hoi viet lai:"
            )
        else:
            prompt = (
                "Based on the conversation history below, rewrite the user's latest question into a standalone query "
                "with complete subject, predicate, and document-retrieval keywords. "
                "Return ONLY the rewritten question with no extra explanation.\n\n"
                f"Conversation history:\n{os.linesep.join(history_lines)}\n\n"
                f"Original question: {question}\n"
                "Rewritten question:"
            )

        response = self._ollama_client().generate(model="qwen2.5:1.5b", prompt=prompt)
        if hasattr(response, "response"):
            rewritten = response.response
        else:
            rewritten = response.get("response", "")
        rewritten = (rewritten or "").strip()
        return rewritten if rewritten else question

    def _build_chat_context(self, docs: List[Document]) -> str:
        sections = []
        for index, doc in enumerate(docs, start=1):
            source_name = doc.metadata.get("source_name", "unknown")
            page_number = doc.metadata.get("page_number") or doc.metadata.get("page") or "N/A"
            sections.append(f"[{index}] source={source_name} page={page_number}\n{doc.page_content}")
        return "\n\n".join(sections)

    def chat(
        self,
        notebook_id: str,
        question: str,
        model_name: str,
        top_k: int,
        session_id: Optional[str],
        source_names: Optional[List[str]],
        answer_language: Optional[str] = None,
        chat_settings: Optional[dict[str, str]] = None,
    ) -> dict:
        print(f"[CHAT] Starting chat - notebook_id={notebook_id}, question={question[:50]}...")
        if not self.storage.notebook_exists(notebook_id):
            print(f"[CHAT ERROR] Notebook not found: {notebook_id}")
            raise ValueError("Notebook not found")

        dirs = self.storage.get_notebook_dirs(notebook_id)
        index_file = dirs["vector_db"] / "index.faiss"
        if not index_file.exists():
            sources = self.storage.list_source_documents(notebook_id)
            if not sources:
                raise ValueError("Notebook has no source documents yet. Upload files first.")

            # Auto-index once so chat can work immediately after upload.
            self.index_notebook(
                notebook_id=notebook_id,
                source_ids=None,
                chunk_size=1000,
                chunk_overlap=100,
            )

            if not index_file.exists():
                raise ValueError("Notebook indexing failed. Please try again.")

        session_id = self.storage.ensure_chat_session(notebook_id=notebook_id, session_id=session_id)
        print(f"[CHAT] Session ID: {session_id}")
        
        rewritten_question = self._rewrite_query(
            notebook_id=notebook_id,
            session_id=session_id,
            question=question,
        )
        print(f"[CHAT] Rewritten question: {rewritten_question}")

        try:
            print(f"[CHAT] Loading vector store from: {dirs['vector_db']}")
            vector_store = FAISS.load_local(
                folder_path=str(dirs["vector_db"]),
                embeddings=self.embedder,
                allow_dangerous_deserialization=True,
            )
            print(f"[CHAT] Vector store loaded successfully")
        except Exception as e:
            print(f"[CHAT ERROR] Failed to load vector store: {str(e)}")
            raise

        metadata_filter = {"notebook_id": notebook_id}
        if source_names:
            metadata_filter["source_name"] = {"$in": source_names}

        retrieval_k = max(top_k, 8 if len(question.split()) <= 8 else 5)
        print(f"[CHAT] Retrieval K: {retrieval_k}, metadata_filter: {metadata_filter}")

        try:
            print(f"[CHAT] Starting similarity search...")
            docs = vector_store.similarity_search(
                query=rewritten_question,
                k=retrieval_k,
                filter=metadata_filter,
            )
            print(f"[CHAT] Similarity search returned {len(docs)} documents")
        except Exception as e:
            print(f"[CHAT ERROR] Similarity search failed: {str(e)}")
            raise

        if not docs:
            print(f"[CHAT] No documents found, returning default message")
            answer = "Khong tim thay ngu canh phu hop trong notebook de tra loi cau hoi nay."
            citations = []
            prompt = ""
        else:
            try:
                print(f"[CHAT] Building context from {len(docs)} documents")
                resolved_language = answer_language or self._detect_answer_language(question)
                print(f"[CHAT] Resolved language: {resolved_language}")
                context = self._build_chat_context(docs)
                print(f"[CHAT] Context built, length: {len(context)}")
                
                prompt_template = self._build_prompt_template(resolved_language, chat_settings)
                prompt = prompt_template.format(context=context, question=rewritten_question)
                print(f"[CHAT] Prompt formatted, length: {len(prompt)}")

                print(f"[CHAT] Calling Ollama with model: {model_name}")
                response = self._ollama_client().generate(
                    model=model_name,
                    prompt=prompt,
                    options={
                        "temperature": 0.1,
                        "top_p": 0.85,
                        "repeat_penalty": 1.05,
                        "num_predict": 220,
                    },
                )
                print(f"[CHAT] Ollama response received")
                
                if hasattr(response, "response"):
                    answer = response.response
                else:
                    answer = response.get("response", "")
                answer = answer.strip()
                print(f"[CHAT] Answer extracted, length: {len(answer)}")

                print(f"[CHAT] Processing citations from {len(docs)} documents")
                citations = []
                for idx, doc in enumerate(docs):
                    try:
                        citation = {
                            "source_name": doc.metadata.get("source_name", "unknown"),
                            "page": doc.metadata.get("page_number") or doc.metadata.get("page"),
                            "snippet": doc.page_content[:280].replace("\n", " "),
                        }
                        citations.append(citation)
                    except Exception as e:
                        print(f"[CHAT ERROR] Failed to process citation {idx}: {str(e)}")
                        raise
                print(f"[CHAT] Processed {len(citations)} citations")
            except Exception as e:
                print(f"[CHAT ERROR] Error in document processing: {str(e)}")
                raise

#gọi graph-rag       
        print(f"[CHAT] Calling GraphRAG...")
        try:
            answer_graph = self.GRS.answer_question(notebook_id, rewritten_question)
            print(f"[CHAT] GraphRAG response received, length: {len(answer_graph) if answer_graph else 0}")
        except Exception as e:
            # Silently fail graph RAG, use empty string as fallback
            import logging
            print(f"[CHAT ERROR] GraphRAG failed: {str(e)}")
            logging.warning(f"GraphRAG failed for notebook {notebook_id}: {str(e)}")
            answer_graph = ""

        try:
            print(f"[CHAT] Saving chat messages to storage...")
            self.storage.add_chat_message(notebook_id=notebook_id, session_id=session_id, role="user", content=question)
            print(f"[CHAT] User message saved")
            self.storage.add_chat_message(notebook_id=notebook_id, session_id=session_id, role="assistant", content=answer)
            print(f"[CHAT] Assistant message saved")
            # Note: answer_graph is returned in API response but NOT saved as separate message
            # Database schema only supports roles: 'system', 'user', 'assistant'
            print(f"[CHAT] Chat messages saved successfully")
        except Exception as e:
            print(f"[CHAT ERROR] Failed to save chat messages: {str(e)}")
            raise

        print(f"[CHAT] Preparing response...")
        return {
            "notebook_id": notebook_id,
            "session_id": session_id,
            "rewritten_question": rewritten_question,
            "answer": answer,
            "answer_graph": answer_graph,
            "citations": citations,
        }
